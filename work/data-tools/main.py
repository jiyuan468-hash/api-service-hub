"""Data Processing API - Document conversion and data extraction service."""

import os
import uuid
import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any

import yaml
import pandas as pd
import pdfplumber
from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# ─── Config ────────────────────────────────────────────────────
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "config.yaml")
with open(CONFIG_PATH, "r", encoding="utf-8") as f:
    config = yaml.safe_load(f)

MAX_FILE_SIZE = config.get("max_file_size_mb", 20) * 1024 * 1024
STORE_DIR = Path(config.get("storage_dir", "./storage"))
STORE_DIR.mkdir(parents=True, exist_ok=True)

# In-memory stats
processing_stats: Dict[str, dict] = {}

# ─── Auth ──────────────────────────────────────────────────────
REQUIRED_API_KEY = config.get("auth", {}).get("api_key", "")

def validate_api_key(x_api_key: Optional[str] = None):
    if REQUIRED_API_KEY and x_api_key != REQUIRED_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API Key")
    return x_api_key

# ─── Pydantic Models ──────────────────────────────────────────
class ConversionResponse(BaseModel):
    job_id: str
    status: str
    result_url: Optional[str] = None
    message: str
    created_at: str = ""

class StatsResponse(BaseModel):
    doc_id: str
    file_name: str
    file_type: str
    file_size_bytes: int
    processed_at: str
    pages: Optional[int] = None

# ─── Validation Helpers ────────────────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".xls"}

async def save_upload(file: UploadFile, tmp_dir: Path) -> Path:
    """Save uploaded file with size check and extension validation."""
    # Check size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // 1024 // 1024}MB)")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    dest = tmp_dir / f"{uuid.uuid4().hex}{ext}"
    dest.write_bytes(content)
    return dest

# ─── Conversion Functions ──────────────────────────────────────
def convert_pdf_to_text(file_path: Path) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {page_num} ---\n{text}")
    return "\n\n".join(text_parts)

def convert_excel_to_csv(file_path: Path) -> bytes:
    """Convert Excel to CSV for each sheet."""
    output = b""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        csv_buffer = df.to_csv(index=False)
        output += f"=== Sheet: {sheet_name} ===\n{csv_buffer}\n\n".encode("utf-8")
    return output

def convert_docx_to_markdown(file_path: Path) -> str:
    """Convert Word document to Markdown."""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        raise ImportError("python-docx required. Install: pip install python-docx")

    doc = Document(str(file_path))
    md_lines = []

    for para in doc.paragraphs:
        style = para.style.name
        if style.startswith("Heading"):
            level = style.replace("Heading ", "").strip()
            try:
                level_num = int(level)
                md_lines.append(f"{'#' * level_num} {para.text}")
            except ValueError:
                md_lines.append(f"## {para.text}")
        elif style == "List":
            md_lines.append(f"- {para.text}")
        elif para.text.strip():
            md_lines.append(para.text)

    # Extract tables
    for table in doc.tables:
        md_lines.append("\n| " + " | ".join(cell.text for cell in table.rows[0].cells) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |")
        for row in table.rows[1:]:
            md_lines.append("| " + " | ".join(cell.text for cell in row.cells) + " |")
        md_lines.append("")

    return "\n".join(md_lines)

# ─── App ───────────────────────────────────────────────────────
app = FastAPI(
    title="Data Processing API",
    description="Document conversion and data extraction service",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/health")
async def health_check(api_key: str = Depends(validate_api_key)):
    return {
        "status": "ok",
        "version": "1.0.0",
        "timestamp": datetime.now().isoformat(),
        "supported_formats": list(ALLOWED_EXTENSIONS),
        "max_file_size_mb": MAX_FILE_SIZE // 1024 // 1024
    }

@app.post("/convert/pdf-to-text", response_model=ConversionResponse)
async def pdf_to_text(
    file: UploadFile = File(...),
    api_key: str = Depends(validate_api_key)
):
    """Extract text content from a PDF file."""
    job_id = uuid.uuid4().hex[:12]

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        src = await save_upload(file, tmp_path)

        try:
            text = convert_pdf_to_text(src)
            # Save result
            result_path = STORE_DIR / f"{job_id}_output.txt"
            result_path.write_text(text, encoding="utf-8")

            processing_stats[job_id] = {
                "doc_id": job_id,
                "file_name": file.filename,
                "file_type": "pdf",
                "file_size_bytes": src.stat().st_size,
                "processed_at": datetime.now().isoformat(),
                "pages": None,
                "result_file": str(result_path)
            }

            return ConversionResponse(
                job_id=job_id,
                status="success",
                result_url=f"/results/{job_id}_output.txt",
                message="PDF converted successfully",
                created_at=datetime.now().isoformat()
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/excel-to-csv", response_model=ConversionResponse)
async def excel_to_csv(
    file: UploadFile = File(...),
    api_key: str = Depends(validate_api_key)
):
    """Convert Excel file to CSV format."""
    job_id = uuid.uuid4().hex[:12]

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        src = await save_upload(file, tmp_path)

        try:
            csv_data = convert_excel_to_csv(src)
            result_path = STORE_DIR / f"{job_id}_output.csv"
            result_path.write_bytes(csv_data)

            processing_stats[job_id] = {
                "doc_id": job_id,
                "file_name": file.filename,
                "file_type": "excel",
                "file_size_bytes": src.stat().st_size,
                "processed_at": datetime.now().isoformat(),
                "result_file": str(result_path)
            }

            return ConversionResponse(
                job_id=job_id,
                status="success",
                result_url=f"/results/{job_id}_output.csv",
                message="Excel converted to CSV successfully",
                created_at=datetime.now().isoformat()
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/docx-to-markdown", response_model=ConversionResponse)
async def docx_to_markdown(
    file: UploadFile = File(...),
    api_key: str = Depends(validate_api_key)
):
    """Convert Word document to Markdown format."""
    job_id = uuid.uuid4().hex[:12]

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        src = await save_upload(file, tmp_path)

        try:
            md_text = convert_docx_to_markdown(src)
            result_path = STORE_DIR / f"{job_id}_output.md"
            result_path.write_text(md_text, encoding="utf-8")

            processing_stats[job_id] = {
                "doc_id": job_id,
                "file_name": file.filename,
                "file_type": "docx",
                "file_size_bytes": src.stat().st_size,
                "processed_at": datetime.now().isoformat(),
                "result_file": str(result_path)
            }

            return ConversionResponse(
                job_id=job_id,
                status="success",
                result_url=f"/results/{job_id}_output.md",
                message="Word document converted to Markdown",
                created_at=datetime.now().isoformat()
            )
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

@app.get("/stats/{doc_id}", response_model=StatsResponse)
async def get_processing_stats(doc_id: str):
    """Get statistics about a previously processed document."""
    stat = processing_stats.get(doc_id)
    if not stat:
        raise HTTPException(status_code=404, detail="Document ID not found")
    return StatsResponse(**stat)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8001, reload=True)
